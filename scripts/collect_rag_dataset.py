#!/usr/bin/env python3
"""
Korean Government RAG Test Dataset Generator

Generates 400 realistic Korean government documents across 4 categories:
- 공문서 (Official Documents): 40 txt, 25 pdf, 20 docx, 15 xlsx
- 메뉴얼 (Manuals): 40 txt, 25 pdf, 20 docx, 15 xlsx
- 공시자료 (Disclosure Materials): 40 txt, 25 pdf, 20 docx, 15 xlsx
- 보도자료 (Press Releases): 40 txt, 25 pdf, 20 docx, 15 xlsx
"""

import json
import os
import random
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# ─────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────
OUTPUT_BASE = Path("/home/siujang/Documents/GovOn/data/raw/rag_test_dataset")
KOREAN_FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
CATEGORIES = ["공문서", "메뉴얼", "공시자료", "보도자료"]
FORMAT_DIST = {"txt": 40, "pdf": 25, "docx": 20, "xlsx": 15}

random.seed(42)


# ─────────────────────────────────────────────
# Shared Data
# ─────────────────────────────────────────────
MINISTRIES = [
    "행정안전부", "기획재정부", "과학기술정보통신부", "보건복지부", "교육부",
    "국토교통부", "환경부", "산업통상자원부", "고용노동부", "외교부",
    "법무부", "국방부", "농림축산식품부", "해양수산부", "문화체육관광부",
    "여성가족부", "중소벤처기업부", "금융위원회", "공정거래위원회", "방송통신위원회",
]
COMPANIES = [
    "삼성전자㈜", "현대자동차㈜", "SK하이닉스㈜", "LG전자㈜", "포스코홀딩스㈜",
    "카카오㈜", "네이버㈜", "셀트리온㈜", "한국전력공사", "한국가스공사",
    "대한항공㈜", "롯데쇼핑㈜", "GS칼텍스㈜", "현대건설㈜", "두산에너빌리티㈜",
    "KB금융지주㈜", "신한금융지주㈜", "하나금융지주㈜", "우리금융지주㈜",
    "한국투자증권㈜",
]
CITIES = ["서울특별시", "부산광역시", "대구광역시", "인천광역시", "광주광역시",
          "대전광역시", "울산광역시", "세종특별자치시", "수원시", "창원시"]

DOC_TYPES = ["행정명령", "공고", "훈령", "예규", "고시", "지시", "협조공문", "통보"]
MANUAL_TYPES = ["업무처리 매뉴얼", "시스템 사용 지침", "민원처리 절차서", "안전관리 지침",
                "개인정보보호 운영지침", "비상대응 절차서", "계약업무 처리요령"]
PRESS_TOPICS = [
    "디지털 전환 가속화 정책", "탄소중립 실현 방안", "스마트시티 구축사업",
    "청년 일자리 창출 대책", "저출산 극복 종합대책", "인공지능 산업 육성",
    "전기차 보급 확대 정책", "수소경제 로드맵", "반도체 클러스터 조성",
    "K-바이오 글로벌 경쟁력 강화", "농업 스마트화 지원", "해양 환경보전 대책",
    "공공데이터 개방 확대", "사이버보안 강화 방안", "국가 R&D 예산 배분 결과",
]


def rand_date(start_year=2022, end_year=2025) -> str:
    start = datetime(start_year, 1, 1)
    end = datetime(end_year, 12, 31)
    delta = (end - start).days
    return (start + timedelta(days=random.randint(0, delta))).strftime("%Y년 %m월 %d일")


def rand_doc_number() -> str:
    ministry_code = "".join([str(random.randint(0, 9)) for _ in range(4)])
    return f"{random.randint(2022, 2025)}-{ministry_code}-{random.randint(1000, 9999)}"


def rand_phone() -> str:
    return f"02-{random.randint(100, 999)}-{random.randint(1000, 9999)}"


def rand_amount(min_b=1, max_b=500) -> str:
    amt = random.randint(min_b * 100, max_b * 100) / 100
    return f"{amt:,.1f}억 원"


# ─────────────────────────────────────────────
# Content Generators
# ─────────────────────────────────────────────

def make_gongmunso_content(idx: int) -> dict:
    """공문서 (Official Document) content"""
    ministry = random.choice(MINISTRIES)
    doc_type = random.choice(DOC_TYPES)
    recipient_ministry = random.choice([m for m in MINISTRIES if m != ministry])
    date = rand_date()
    doc_num = rand_doc_number()
    topics = [
        "디지털정부 서비스 개선 추진",
        "공공데이터 활용 활성화 방안",
        "행정업무 효율화 시스템 도입",
        "공무원 역량강화 교육 실시",
        "정보보안 강화 조치 시행",
        "예산집행 점검 결과 보고",
        "국정감사 자료 제출 요청",
        "부처 간 협업사업 추진 계획",
        "민원처리 개선 대책 시행",
        "공공기관 경영평가 결과 통보",
    ]
    topic = random.choice(topics)
    officer_name = random.choice(["김민준", "이서연", "박지훈", "최유진", "정현우",
                                   "강수빈", "윤도현", "임하은", "조성민", "신지아"])
    dept = random.choice(["행정관리담당관실", "정책기획과", "운영지원과", "디지털정부국", "재정관리과"])

    body_paragraphs = [
        f"1. 관련 근거\n  가. 「전자정부법」 제36조 및 같은 법 시행령 제39조\n  나. 행정업무의 운영 및 혁신에 관한 규정(대통령령) 제3조\n  다. {ministry} 업무처리 지침 제{random.randint(5, 30)}조",
        f"\n2. 추진 배경\n  {topic}과 관련하여 {date}부터 본 부처에서 시행 중인 정책의 효과성을 제고하고, 관련 기관과의 긴밀한 협조를 위해 아래와 같이 통보합니다.",
        f"\n3. 세부 추진 사항\n  가. 사업명: {topic}\n  나. 추진 기간: {rand_date()} ~ {rand_date()}\n  다. 소요 예산: {rand_amount()}\n  라. 주관 부서: {ministry} {dept}\n  마. 협조 기관: {recipient_ministry}",
        f"\n4. 협조 요청 사항\n  위 사업의 원활한 추진을 위해 귀 기관의 적극적인 협조를 요청드립니다.\n  - 관련 자료 제출 기한: {rand_date()}\n  - 담당자 회의 참석 여부 회신: {rand_date()}까지",
        f"\n붙임: 1. {topic} 추진 계획서 1부\n      2. 협조 요청 사항 목록 1부. 끝.",
    ]

    return {
        "문서번호": f"{ministry} {rand_doc_number()}",
        "수신": recipient_ministry,
        "경유": "(경유 없음)" if random.random() > 0.3 else random.choice(MINISTRIES),
        "제목": topic,
        "내용": "\n".join(body_paragraphs),
        "처리과": dept,
        "담당자": officer_name,
        "연락처": rand_phone(),
        "생산일자": date,
        "문서구분": doc_type,
    }


def make_manual_content(idx: int) -> dict:
    """메뉴얼 (Manual) content"""
    ministry = random.choice(MINISTRIES)
    manual_type = random.choice(MANUAL_TYPES)
    version = f"v{random.randint(1, 5)}.{random.randint(0, 9)}"
    date = rand_date()

    systems = ["전자결재시스템", "민원24", "나라장터", "온-나라 문서시스템",
               "국가재정시스템(dBrain)", "인사혁신처 인사통합시스템", "행정정보공동이용시스템"]
    system = random.choice(systems)

    chapters = [
        f"제1장 총칙\n  제1조(목적) 이 매뉴얼은 {ministry}에서 {system}을(를) 활용한 {manual_type}의 절차와 방법을 규정함으로써 업무의 효율성과 일관성을 확보함을 목적으로 한다.\n  제2조(적용 범위) 이 매뉴얼은 {ministry} 소속 공무원 및 관련 기관 담당자에게 적용된다.\n  제3조(용어의 정의) ① \"{system}\"이란 ...(이하 시스템 설명 생략).",
        f"\n제2장 업무 처리 절차\n  제4조(신청 및 접수)\n    ① 담당자는 해당 업무 발생 시 {system}에 접속하여 관련 화면을 열람한다.\n    ② 접수된 건은 D+{random.randint(1, 5)}일 이내에 처리하여야 한다.\n    ③ 처리 결과는 신청인에게 문자 또는 전자우편으로 통보한다.\n  제5조(검토 및 결재)\n    ① 담당자 → 팀장 → 과장 → 국장 순으로 결재를 진행한다.\n    ② 긴급 사안의 경우 전결 규정에 따라 처리한다.",
        f"\n제3장 주요 기능 안내\n  3.1 로그인 및 권한 설정\n    - 접속 URL: https://system.go.kr\n    - ID/PW는 행정전산망 계정과 동일하게 사용\n    - 처음 로그인 시 비밀번호 반드시 변경\n  3.2 주요 메뉴 설명\n    [결재함] → [업무처리] → [통계·현황] → [환경설정]\n  3.3 자주 묻는 질문(FAQ)\n    Q: 비밀번호를 잊어버린 경우?\n    A: 관리자(☎ {rand_phone()})에게 연락하여 초기화 요청",
        f"\n제4장 주의사항 및 오류 대응\n  ① 개인정보가 포함된 자료는 반드시 암호화 후 전송한다.\n  ② 시스템 오류 발생 시 오류 화면을 캡처하여 헬프데스크에 접수한다.\n  ③ 민감 정보는 업무 종료 후 즉시 삭제 또는 반납한다.\n  ④ 무단 접근 및 비인가 데이터 열람은 「개인정보 보호법」 위반에 해당한다.",
        f"\n제5장 관련 법령 및 참고자료\n  - 「전자정부법」(법률 제19010호)\n  - 「개인정보 보호법」(법률 제19234호)\n  - {ministry} 내부지침 제{random.randint(10, 100)}호\n  - 행정안전부 고시 제{random.randint(2020, 2024)}-{random.randint(1, 50)}호\n\n부록\n  [별표 1] 업무 처리 흐름도\n  [별표 2] 관련 서식 목록\n  [별표 3] 담당자 연락처",
    ]

    toc = [
        "목차",
        "제1장 총칙 ···················· 3",
        "제2장 업무 처리 절차 ·········· 7",
        "제3장 주요 기능 안내 ·········· 12",
        "제4장 주의사항 및 오류 대응 ··· 18",
        "제5장 관련 법령 및 참고자료 ·· 22",
        "부록 ·························· 25",
    ]

    return {
        "제목": f"{ministry} {manual_type}",
        "버전": version,
        "발행기관": ministry,
        "발행일": date,
        "목차": "\n".join(toc),
        "본문": "\n".join(chapters),
        "비고": f"본 매뉴얼에 대한 문의는 {ministry} 운영지원과(☎ {rand_phone()})로 연락바랍니다.",
    }


def make_gongsi_content(idx: int) -> dict:
    """공시자료 (Disclosure Materials) content"""
    company = random.choice(COMPANIES)
    year = random.randint(2021, 2024)
    quarter = random.randint(1, 4)
    report_type = random.choice(["사업보고서", "반기보고서", "분기보고서", "감사보고서"])
    date = rand_date(year, year)

    # Financial figures
    revenue = random.randint(10000, 500000)
    op_income = int(revenue * random.uniform(0.03, 0.20))
    net_income = int(op_income * random.uniform(0.6, 0.95))
    assets = int(revenue * random.uniform(1.2, 3.0))
    liabilities = int(assets * random.uniform(0.3, 0.65))
    equity = assets - liabilities
    eps = random.randint(500, 30000)

    sectors = ["반도체", "전자기기", "자동차", "화학", "금융서비스", "IT서비스",
               "바이오헬스케어", "에너지", "건설", "유통", "통신서비스"]
    sector = random.choice(sectors)

    body = f"""I. 회사의 개요
1. 회사의 명칭: {company}
2. 사업의 내용: {sector} 관련 제품 및 서비스 제공
3. 본점 소재지: {random.choice(CITIES)} {random.choice(['강남구', '중구', '서초구', '영등포구', '종로구'])}
4. 설립연월일: {random.randint(1960, 2010)}년 {random.randint(1, 12)}월 {random.randint(1, 28)}일
5. 결산월: {random.randint(10, 12)}월

II. 주요 사업 내용
  {company}은(는) {sector} 분야의 선도 기업으로, 국내외 시장에서 고객 가치 창출을 위해 지속적인 연구개발과 품질혁신을 추진하고 있습니다.
  주요 제품군은 다음과 같습니다:
  - 제품/서비스 A: 매출 비중 {random.randint(30, 60)}%
  - 제품/서비스 B: 매출 비중 {random.randint(15, 30)}%
  - 제품/서비스 C: 매출 비중 {random.randint(10, 20)}%
  - 기타: 나머지

III. 재무제표 요약 (단위: 억 원)
  ┌─────────────────────┬─────────────────┐
  │ 항목                 │ 금액             │
  ├─────────────────────┼─────────────────┤
  │ 매출액               │ {revenue:,}       │
  │ 영업이익             │ {op_income:,}     │
  │ 당기순이익           │ {net_income:,}    │
  │ 자산총계             │ {assets:,}        │
  │ 부채총계             │ {liabilities:,}   │
  │ 자본총계             │ {equity:,}        │
  │ 주당순이익(EPS)      │ {eps:,}원         │
  └─────────────────────┴─────────────────┘
  * 영업이익률: {op_income/revenue*100:.1f}%
  * 부채비율: {liabilities/equity*100:.1f}%

IV. 감사의견
  {'적정' if random.random() > 0.05 else '한정'} (감사법인: {random.choice(['삼일회계법인', '삼정KPMG', '한영회계법인', 'EY한영', 'BDO코리아'])})

V. 주요 위험 요소
  1. 시장 위험: 글로벌 경기 침체 및 환율 변동에 따른 매출 영향
  2. 운영 위험: 공급망 불안정 및 원자재 가격 상승
  3. 규제 위험: 국내외 환경·안전 규제 강화
  4. 경쟁 위험: 동종 업계 경쟁 심화 및 신규 진입자 출현

VI. 이사회 현황
  - 사내이사 {random.randint(2, 5)}명, 사외이사 {random.randint(3, 7)}명
  - 이사회 개최 횟수(당기): {random.randint(5, 15)}회"""

    return {
        "회사명": company,
        "보고서종류": report_type,
        "사업연도": f"{year}년 {quarter}분기",
        "제출일": date,
        "본문": body,
        "매출액": f"{revenue:,}억 원",
        "영업이익": f"{op_income:,}억 원",
        "당기순이익": f"{net_income:,}억 원",
    }


def make_bodo_content(idx: int) -> dict:
    """보도자료 (Press Release) content"""
    ministry = random.choice(MINISTRIES)
    topic = random.choice(PRESS_TOPICS)
    date = rand_date()
    officer = random.choice(["홍길동", "김영희", "이철수", "박민지", "최동욱",
                              "정수아", "강병호", "윤지현", "임태호", "조미래"])
    dept = random.choice(["대변인실", "홍보담당관실", "정책기획과", "소통홍보팀"])
    budget = rand_amount(50, 10000)
    target_year = random.randint(2024, 2030)

    headlines = [
        f"{ministry}, '{topic}' 본격 시동",
        f"정부, {topic} 위해 {budget} 투입 결정",
        f"{ministry} '{topic}' 추진 성과 발표",
        f"2024년 {topic} 주요 성과 및 향후 계획",
        f"{topic} 관련 새 정책 시행… {target_year}년까지 {budget} 지원",
    ]
    headline = random.choice(headlines)

    achievements = [
        f"참여 기업·기관 {random.randint(100, 5000)}개소 달성",
        f"직접 수혜 국민 {random.randint(10, 500)}만 명 돌파",
        f"예산 집행률 {random.randint(85, 99)}% 기록",
        f"관련 일자리 {random.randint(5000, 100000):,}개 창출",
        f"온실가스 감축 {random.randint(10, 500)}만 톤CO₂eq 달성",
    ]

    body = f"""□ {ministry}은(는) {date} {topic}과 관련한 주요 성과와 향후 추진 계획을 발표하였다.

□ 주요 내용
  ○ {ministry}은(는) 지난 {random.randint(1, 3)}년간 {topic}을(를) 적극 추진하여 다음과 같은 성과를 거두었다.
    - {random.choice(achievements)}
    - {random.choice(achievements)}
    - {random.choice(achievements)}

□ 향후 계획
  ○ {target_year}년까지 총 {budget}을 투입하여 사업을 지속 확대할 예정이다.
  ○ 세부 추진 전략은 다음과 같다.
    - 1단계({random.randint(2024, 2025)}년): 기반 구축 및 시범사업 추진
    - 2단계({random.randint(2025, 2027)}년): 본사업 확대 및 성과 점검
    - 3단계({random.randint(2027, 2030)}년): 전국 확산 및 고도화

□ 기대 효과
  ○ {topic} 추진을 통해 국민의 삶의 질 향상 및 국가 경쟁력 강화에 기여할 것으로 기대된다.
  ○ 경제적 파급 효과는 연간 {rand_amount(1000, 50000)}으로 추산된다.

□ {ministry} 장관은 "{topic}은 우리 정부의 핵심 국정과제로, 현장 중심의 정책 추진을 통해 국민이 체감하는 성과를 만들어 나가겠다"고 밝혔다.

※ 자세한 내용은 첨부 자료를 참고하시기 바랍니다.

【문의처】 {ministry} {dept} 담당자 {officer} (☎ {rand_phone()})"""

    return {
        "발표기관": ministry,
        "담당부서": dept,
        "담당자": officer,
        "연락처": rand_phone(),
        "발표일": date,
        "제목": headline,
        "본문": body,
    }


CONTENT_MAKERS = {
    "공문서": make_gongmunso_content,
    "메뉴얼": make_manual_content,
    "공시자료": make_gongsi_content,
    "보도자료": make_bodo_content,
}


def content_to_text(category: str, data: dict) -> str:
    """Convert content dict to formatted text"""
    lines = []
    if category == "공문서":
        lines += [
            "=" * 60,
            f"【 행정 공문서 】  문서구분: {data['문서구분']}",
            "=" * 60,
            f"문서번호: {data['문서번호']}",
            f"수  신: {data['수신']}",
            f"경  유: {data['경유']}",
            f"생산일자: {data['생산일자']}",
            f"제  목: {data['제목']}",
            "",
            data["내용"],
            "",
            f"처리과: {data['처리과']}",
            f"담당자: {data['담당자']}  연락처: {data['연락처']}",
            "=" * 60,
        ]
    elif category == "메뉴얼":
        lines += [
            "=" * 60,
            f"【 {data['제목']} 】",
            f"발행기관: {data['발행기관']}  버전: {data['버전']}  발행일: {data['발행일']}",
            "=" * 60,
            "",
            data["목차"],
            "",
            data["본문"],
            "",
            f"※ {data['비고']}",
        ]
    elif category == "공시자료":
        lines += [
            "=" * 60,
            f"【 {data['보고서종류']} 】",
            f"회사명: {data['회사명']}  사업연도: {data['사업연도']}  제출일: {data['제출일']}",
            "=" * 60,
            "",
            data["본문"],
        ]
    elif category == "보도자료":
        lines += [
            "=" * 60,
            "【 보 도 자 료 】",
            f"발표기관: {data['발표기관']}  |  발표일: {data['발표일']}",
            "=" * 60,
            f"제목: {data['제목']}",
            "",
            data["본문"],
        ]
    return "\n".join(lines)


# ─────────────────────────────────────────────
# File Writers
# ─────────────────────────────────────────────

def write_txt(path: Path, category: str, data: dict):
    text = content_to_text(category, data)
    path.write_text(text, encoding="utf-8")


def _safe_pdf_text(text: str) -> str:
    """Sanitise a string for fpdf2: replace box-drawing and other problematic characters."""
    # str.maketrans only accepts single-char keys
    BOX_CHARS = str.maketrans({
        "┌": "+", "┐": "+", "└": "+", "┘": "+",
        "├": "+", "┤": "+", "┬": "+", "┴": "+", "┼": "+",
        "─": "-", "│": "|",
        "=": "-",           # '=' can overflow in CJK fonts
        "★": "*", "☆": "*",
        "①": "1", "②": "2", "③": "3", "④": "4", "⑤": "5",
        "₂": "2",
    })
    # Also remove any characters outside the BMP that the font may not have
    result = []
    for ch in text.translate(BOX_CHARS):
        cp = ord(ch)
        if cp > 0xFFFF:
            result.append("?")
        else:
            result.append(ch)
    return "".join(result)


def write_pdf(path: Path, category: str, data: dict):
    # Build structured content instead of raw text
    if category == "공문서":
        title = f"행정공문서: {data['제목']}"
        sections = [
            f"문서번호: {data['문서번호']}",
            f"수신: {data['수신']}  /  경유: {data['경유']}",
            f"생산일자: {data['생산일자']}  /  문서구분: {data['문서구분']}",
            "",
            data["내용"],
            "",
            f"처리과: {data['처리과']}  /  담당자: {data['담당자']}  /  연락처: {data['연락처']}",
        ]
    elif category == "메뉴얼":
        title = data["제목"]
        sections = [
            f"발행기관: {data['발행기관']}  /  버전: {data['버전']}  /  발행일: {data['발행일']}",
            "",
            data["목차"],
            "",
            data["본문"],
            "",
            data["비고"],
        ]
    elif category == "공시자료":
        title = f"{data['회사명']} {data['보고서종류']}"
        sections = [
            f"사업연도: {data['사업연도']}  /  제출일: {data['제출일']}",
            f"매출액: {data['매출액']}  /  영업이익: {data['영업이익']}  /  당기순이익: {data['당기순이익']}",
            "",
            data["본문"],
        ]
    else:  # 보도자료
        title = data["제목"]
        sections = [
            f"발표기관: {data['발표기관']}  /  발표일: {data['발표일']}",
            f"담당부서: {data['담당부서']}  /  담당자: {data['담당자']}  /  연락처: {data['연락처']}",
            "",
            data["본문"],
        ]

    full_text = "\n".join(sections)

    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    pdf.add_font("NotoSansCJK", "", KOREAN_FONT_PATH)

    # Title
    pdf.set_font("NotoSansCJK", size=13)
    pdf.set_fill_color(220, 230, 242)
    safe_title = _safe_pdf_text(title[:50])
    pdf.cell(0, 12, text=safe_title, fill=True, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Body — split on newlines and render each line safely
    pdf.set_font("NotoSansCJK", size=9)
    effective_width = pdf.w - pdf.l_margin - pdf.r_margin
    # Estimate ~1.8 chars per mm for 9pt CJK font (conservative)
    max_chars = int(effective_width / 1.8)

    for raw_line in full_text.split("\n"):
        safe_line = _safe_pdf_text(raw_line)
        if not safe_line.strip():
            pdf.ln(4)
            continue
        # Break into chunks that fit the page width
        while len(safe_line) > max_chars:
            chunk = safe_line[:max_chars]
            safe_line = safe_line[max_chars:]
            pdf.cell(0, 5, text=chunk, new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 5, text=safe_line, new_x="LMARGIN", new_y="NEXT")

    pdf.output(str(path))


def write_docx(path: Path, category: str, data: dict):
    doc = Document()

    # Title
    if category == "공문서":
        title_text = f"행정공문서: {data['제목']}"
    elif category == "메뉴얼":
        title_text = data["제목"]
    elif category == "공시자료":
        title_text = f"{data['회사명']} {data['보고서종류']}"
    else:
        title_text = data["제목"]

    title_para = doc.add_heading(title_text, level=1)
    title_para.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Metadata table
    doc.add_paragraph()
    if category == "공문서":
        meta_rows = [
            ("문서번호", data["문서번호"]),
            ("수신", data["수신"]),
            ("생산일자", data["생산일자"]),
            ("처리과", data["처리과"]),
            ("담당자", f"{data['담당자']} (☎ {data['연락처']})"),
        ]
    elif category == "메뉴얼":
        meta_rows = [
            ("발행기관", data["발행기관"]),
            ("버전", data["버전"]),
            ("발행일", data["발행일"]),
        ]
    elif category == "공시자료":
        meta_rows = [
            ("회사명", data["회사명"]),
            ("보고서종류", data["보고서종류"]),
            ("사업연도", data["사업연도"]),
            ("제출일", data["제출일"]),
            ("매출액", data["매출액"]),
        ]
    else:
        meta_rows = [
            ("발표기관", data["발표기관"]),
            ("담당부서", data["담당부서"]),
            ("담당자", f"{data['담당자']} (☎ {data['연락처']})"),
            ("발표일", data["발표일"]),
        ]

    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(meta_rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # Body content
    text = content_to_text(category, data)
    for line in text.split("\n"):
        if line.startswith("제") and "장" in line[:5]:
            p = doc.add_heading(line, level=2)
        elif line.startswith("  제") and "조" in line[:7]:
            p = doc.add_heading(line.strip(), level=3)
        elif line.startswith("=" * 5):
            doc.add_paragraph()
        else:
            doc.add_paragraph(line)

    doc.save(str(path))


def write_xlsx(path: Path, category: str, idx: int):
    wb = openpyxl.Workbook()
    ws = wb.active

    header_fill = PatternFill(start_color="1F497D", end_color="1F497D", fill_type="solid")
    header_font = Font(name="맑은 고딕", bold=True, color="FFFFFF", size=11)
    alt_fill = PatternFill(start_color="DCE6F1", end_color="DCE6F1", fill_type="solid")
    body_font = Font(name="맑은 고딕", size=10)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    thin = Side(border_style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def style_header(cell):
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_align
        cell.border = border

    def style_cell(cell, alt=False):
        if alt:
            cell.fill = alt_fill
        cell.font = body_font
        cell.alignment = left_align
        cell.border = border

    if category == "공문서":
        ws.title = "공문서 목록"
        headers = ["번호", "문서번호", "문서구분", "발신기관", "수신기관", "제목", "생산일자", "담당자", "연락처"]
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 22
        ws.column_dimensions["C"].width = 12
        ws.column_dimensions["D"].width = 16
        ws.column_dimensions["E"].width = 16
        ws.column_dimensions["F"].width = 30
        ws.column_dimensions["G"].width = 14
        ws.column_dimensions["H"].width = 10
        ws.column_dimensions["I"].width = 16

        for col, h in enumerate(headers, 1):
            style_header(ws.cell(1, col, h))

        doc_types_list = DOC_TYPES
        for r in range(2, 52):
            ministry = random.choice(MINISTRIES)
            recipient = random.choice([m for m in MINISTRIES if m != ministry])
            topics_short = ["디지털정부 서비스", "공공데이터 활용", "행정업무 효율화", "공무원 교육",
                            "정보보안 조치", "예산집행 점검", "민원처리 개선", "협업사업 추진"]
            row_data = [
                r - 1,
                rand_doc_number(),
                random.choice(doc_types_list),
                ministry,
                recipient,
                random.choice(topics_short) + " 관련 협조 요청",
                rand_date(),
                random.choice(["김민준", "이서연", "박지훈", "최유진"]),
                rand_phone(),
            ]
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(r, col, val)
                style_cell(cell, alt=(r % 2 == 0))

    elif category == "메뉴얼":
        ws.title = "매뉴얼 관리 현황"
        headers = ["번호", "매뉴얼명", "발행기관", "버전", "발행일", "적용대상", "페이지수", "검토주기", "최종검토일"]
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 16
        ws.column_dimensions["D"].width = 8
        ws.column_dimensions["E"].width = 14
        ws.column_dimensions["F"].width = 16
        ws.column_dimensions["G"].width = 10
        ws.column_dimensions["H"].width = 12
        ws.column_dimensions["I"].width = 14

        for col, h in enumerate(headers, 1):
            style_header(ws.cell(1, col, h))

        for r in range(2, 52):
            row_data = [
                r - 1,
                random.choice(MANUAL_TYPES),
                random.choice(MINISTRIES),
                f"v{random.randint(1, 5)}.{random.randint(0, 9)}",
                rand_date(),
                random.choice(["전체 공무원", "담당자", "관리자급 이상", "계약담당자"]),
                random.randint(20, 200),
                random.choice(["연 1회", "2년마다", "수시"]),
                rand_date(),
            ]
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(r, col, val)
                style_cell(cell, alt=(r % 2 == 0))

    elif category == "공시자료":
        ws.title = "재무현황"
        headers = ["번호", "회사명", "보고서종류", "사업연도", "매출액(억)", "영업이익(억)",
                   "당기순이익(억)", "자산총계(억)", "영업이익률(%)", "EPS(원)"]
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 18
        ws.column_dimensions["C"].width = 14
        ws.column_dimensions["D"].width = 14
        for c in ["E", "F", "G", "H"]:
            ws.column_dimensions[c].width = 14
        ws.column_dimensions["I"].width = 14
        ws.column_dimensions["J"].width = 12

        for col, h in enumerate(headers, 1):
            style_header(ws.cell(1, col, h))

        for r in range(2, 52):
            rev = random.randint(5000, 300000)
            op = int(rev * random.uniform(0.03, 0.20))
            net = int(op * random.uniform(0.6, 0.95))
            assets = int(rev * random.uniform(1.2, 3.0))
            row_data = [
                r - 1,
                random.choice(COMPANIES),
                random.choice(["사업보고서", "반기보고서", "분기보고서"]),
                f"{random.randint(2021, 2024)}년 {random.randint(1, 4)}분기",
                rev,
                op,
                net,
                assets,
                round(op / rev * 100, 1),
                random.randint(500, 30000),
            ]
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(r, col, val)
                style_cell(cell, alt=(r % 2 == 0))
                if col in (5, 6, 7, 8, 10) and isinstance(val, (int, float)):
                    cell.number_format = "#,##0"

    else:  # 보도자료
        ws.title = "보도자료 목록"
        headers = ["번호", "발표기관", "제목", "발표일", "담당부서", "담당자",
                   "연락처", "주요내용 요약", "파일형식"]
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 16
        ws.column_dimensions["C"].width = 35
        ws.column_dimensions["D"].width = 14
        ws.column_dimensions["E"].width = 16
        ws.column_dimensions["F"].width = 10
        ws.column_dimensions["G"].width = 16
        ws.column_dimensions["H"].width = 35
        ws.column_dimensions["I"].width = 10

        for col, h in enumerate(headers, 1):
            style_header(ws.cell(1, col, h))

        summaries = [
            "정책 추진 성과 발표 및 향후 계획 설명",
            "신규 사업 시행 공고 및 지원 대상 안내",
            "통계 조사 결과 발표",
            "예산 투입 계획 및 기대 효과 설명",
            "국제 협력 성과 공유 및 MOU 체결 내용",
        ]
        for r in range(2, 52):
            ministry = random.choice(MINISTRIES)
            topic = random.choice(PRESS_TOPICS)
            row_data = [
                r - 1,
                ministry,
                f"{ministry} '{topic}' 관련 발표",
                rand_date(),
                random.choice(["대변인실", "홍보담당관실", "정책기획과"]),
                random.choice(["홍길동", "김영희", "이철수", "박민지"]),
                rand_phone(),
                random.choice(summaries),
                random.choice(["PDF", "HWP", "DOCX"]),
            ]
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(r, col, val)
                style_cell(cell, alt=(r % 2 == 0))

    # Freeze top row
    ws.freeze_panes = "A2"
    ws.row_dimensions[1].height = 20

    wb.save(str(path))


# ─────────────────────────────────────────────
# Main Generator
# ─────────────────────────────────────────────

def generate_all():
    manifest_entries = []
    total_count = 0

    for category in CATEGORIES:
        cat_dir = OUTPUT_BASE / category
        cat_dir.mkdir(parents=True, exist_ok=True)
        print(f"\n[{category}] 생성 시작...")

        idx = 0

        # TXT files
        for i in range(FORMAT_DIST["txt"]):
            data = CONTENT_MAKERS[category](idx)
            filename = f"{category}_{idx+1:03d}.txt"
            fpath = cat_dir / filename
            write_txt(fpath, category, data)
            manifest_entries.append({
                "category": category, "format": "txt", "filename": filename,
                "path": str(fpath.relative_to(OUTPUT_BASE)),
                "created_at": datetime.now().isoformat(),
            })
            idx += 1
        print(f"  ✓ TXT {FORMAT_DIST['txt']}개 완료")

        # PDF files
        for i in range(FORMAT_DIST["pdf"]):
            data = CONTENT_MAKERS[category](idx)
            filename = f"{category}_{idx+1:03d}.pdf"
            fpath = cat_dir / filename
            write_pdf(fpath, category, data)
            manifest_entries.append({
                "category": category, "format": "pdf", "filename": filename,
                "path": str(fpath.relative_to(OUTPUT_BASE)),
                "created_at": datetime.now().isoformat(),
            })
            idx += 1
        print(f"  ✓ PDF {FORMAT_DIST['pdf']}개 완료")

        # DOCX files
        for i in range(FORMAT_DIST["docx"]):
            data = CONTENT_MAKERS[category](idx)
            filename = f"{category}_{idx+1:03d}.docx"
            fpath = cat_dir / filename
            write_docx(fpath, category, data)
            manifest_entries.append({
                "category": category, "format": "docx", "filename": filename,
                "path": str(fpath.relative_to(OUTPUT_BASE)),
                "created_at": datetime.now().isoformat(),
            })
            idx += 1
        print(f"  ✓ DOCX {FORMAT_DIST['docx']}개 완료")

        # XLSX files
        for i in range(FORMAT_DIST["xlsx"]):
            filename = f"{category}_{idx+1:03d}.xlsx"
            fpath = cat_dir / filename
            write_xlsx(fpath, category, idx)
            manifest_entries.append({
                "category": category, "format": "xlsx", "filename": filename,
                "path": str(fpath.relative_to(OUTPUT_BASE)),
                "created_at": datetime.now().isoformat(),
            })
            idx += 1
        print(f"  ✓ XLSX {FORMAT_DIST['xlsx']}개 완료")

        total_count += idx
        print(f"  [{category}] 총 {idx}개 파일 생성 완료")

    # manifest.json
    manifest = {
        "dataset_name": "Korean Government RAG Test Dataset",
        "version": "1.0.0",
        "created_at": datetime.now().isoformat(),
        "total_files": total_count,
        "categories": {
            cat: {
                "total": sum(1 for e in manifest_entries if e["category"] == cat),
                "by_format": {
                    fmt: sum(1 for e in manifest_entries if e["category"] == cat and e["format"] == fmt)
                    for fmt in ["txt", "pdf", "docx", "xlsx"]
                },
            }
            for cat in CATEGORIES
        },
        "format_distribution": FORMAT_DIST,
        "files": manifest_entries,
    }

    manifest_path = OUTPUT_BASE / "manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"\n✓ manifest.json 저장 완료: {manifest_path}")

    # README.md
    readme_lines = [
        "# Korean Government RAG Test Dataset",
        "",
        "한국 정부 문서 기반 RAG(Retrieval-Augmented Generation) 테스트용 데이터셋입니다.",
        "",
        "## 개요",
        "",
        f"- **총 파일 수**: {total_count}개",
        "- **카테고리 수**: 4개",
        "- **생성일**: " + datetime.now().strftime("%Y-%m-%d"),
        "",
        "## 카테고리별 구성",
        "",
        "| 카테고리 | 설명 | TXT | PDF | DOCX | XLSX | 합계 |",
        "|----------|------|-----|-----|------|------|------|",
        "| 공문서 | 행정명령, 공고, 훈령, 예규, 고시 등 | 40 | 25 | 20 | 15 | 100 |",
        "| 메뉴얼 | 업무처리 매뉴얼, 시스템 사용 지침, 민원처리 절차서 | 40 | 25 | 20 | 15 | 100 |",
        "| 공시자료 | 사업보고서, 감사보고서, 분기보고서 | 40 | 25 | 20 | 15 | 100 |",
        "| 보도자료 | 정부 정책 발표, 통계 발표, 사업 성과 보고 | 40 | 25 | 20 | 15 | 100 |",
        "| **합계** | | **160** | **100** | **80** | **60** | **400** |",
        "",
        "## 파일 형식별 특징",
        "",
        "- **TXT**: UTF-8 인코딩, 구조화된 한국어 텍스트",
        "- **PDF**: NotoSansCJK 폰트 사용, 한국어 완전 지원",
        "- **DOCX**: python-docx 생성, 표/제목 스타일 포함",
        "- **XLSX**: openpyxl 생성, 헤더 스타일링 및 데이터 50행 포함",
        "",
        "## 디렉토리 구조",
        "",
        "```",
        "rag_test_dataset/",
        "├── 공문서/          # 행정 공문서 100개",
        "├── 메뉴얼/          # 업무 매뉴얼 100개",
        "├── 공시자료/        # 공시 자료 100개",
        "├── 보도자료/        # 보도 자료 100개",
        "├── manifest.json    # 파일 목록 및 메타데이터",
        "└── README.md        # 이 파일",
        "```",
        "",
        "## 사용 목적",
        "",
        "본 데이터셋은 한국어 정부 문서를 대상으로 하는 RAG 시스템의 성능 평가 및 테스트를 위해 생성된 합성 데이터입니다.",
        "실제 정부 기관의 공식 문서가 아니며, 테스트 및 개발 목적으로만 사용하십시오.",
        "",
        "## 주의사항",
        "",
        "- 본 데이터셋은 자동 생성된 합성 데이터입니다.",
        "- 실제 기관명, 인명, 연락처 등은 테스트 목적으로만 사용된 것입니다.",
        "- 상업적 사용 시 관련 법령을 확인하십시오.",
    ]

    readme_path = OUTPUT_BASE / "README.md"
    readme_path.write_text("\n".join(readme_lines), encoding="utf-8")
    print(f"✓ README.md 저장 완료: {readme_path}")

    # Summary
    print("\n" + "=" * 60)
    print("데이터셋 생성 완료 요약")
    print("=" * 60)
    for cat in CATEGORIES:
        cat_files = [e for e in manifest_entries if e["category"] == cat]
        by_fmt = {fmt: sum(1 for e in cat_files if e["format"] == fmt) for fmt in ["txt", "pdf", "docx", "xlsx"]}
        print(f"  {cat}: {len(cat_files)}개  (txt:{by_fmt['txt']}, pdf:{by_fmt['pdf']}, docx:{by_fmt['docx']}, xlsx:{by_fmt['xlsx']})")
    print(f"\n  총계: {total_count}개")
    print(f"  저장 위치: {OUTPUT_BASE}")
    print("=" * 60)


if __name__ == "__main__":
    generate_all()
